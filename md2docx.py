import markdown
from docx import Document
from bs4 import BeautifulSoup


def md_to_docx(md_file, docx_file):
    if md_file.endswith(".md"):
        # 读取Markdown文件
        with open(md_file, 'r', encoding='utf-8') as f:
            md_text = f.read()
    else:
        md_text = md_file

    html = markdown.markdown(md_text)
    soup = BeautifulSoup(html, 'html.parser')

    # 检查是否存在<body>标签，如果没有，则直接处理soup对象
    if soup.body:
        elements = soup.body.contents
    else:
        elements = soup.contents
    # print(soup)

    # 创建一个新的Word文档
    doc = Document()

    # 遍历BeautifulSoup对象中的段落
    for element in elements:
        if element.name == 'p':
            # 添加段落
            # doc.add_paragraph(element.get_text())
            add_paragraph_bytype(doc, element)
        elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            # 添加标题，根据h标签的级别
            level = int(element.name[1]) - 1
            doc.add_heading(element.get_text(), level=level)
        elif element.name == 'ul':
            # 处理无序列表
            for li in element.find_all('li'):
                # print(li.contents)
                # doc.add_paragraph(li.get_text(), style='ListBullet')
                add_paragraph_bytype(doc, li, style='ListBullet')
        elif element.name == 'ol':
            # 处理有序列表
            for li in element.find_all('li'):
                # print(li.contents)
                doc.add_paragraph(li.get_text(), style='ListBullet')

    new_font_name = '微软雅黑'

    # 遍历文档中的每个段落和运行
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = new_font_name

    # 保存Word文档
    doc.save(docx_file)


def add_paragraph_bytype(doc, element, style=None):
    # 创建段落
    paragraph = doc.add_paragraph(style=style)
    # 遍历元素中的子节点
    for content in element.contents:
        if content.name == 'strong':
            # 处理加粗文本
            run = paragraph.add_run(content.get_text())
            run.bold = True
        elif content.name is None:
            # 处理普通文本
            run = paragraph.add_run(content)
            # print(run.font.name)
        else:
            # 这里可以添加对其他标签的处理逻辑
            pass


# 调用函数，将Markdown文件转换为Word文档
md_to_docx('招标书.md', '招标书.docx')